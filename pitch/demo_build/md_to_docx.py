"""Convert STJP_RESEARCH_REPORT.md to .docx (python-docx).

Handles the report's markdown subset: #/##/### headings, paragraphs with
**bold** / *italic* / `code` inline runs, pipe tables, fenced code blocks,
bullet and numbered lists, horizontal rules.

Usage: python md_to_docx.py <in.md> <out.docx>
"""
import re
import sys

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

ACCENT = RGBColor(0x1F, 0x4E, 0x79)
CODE_BG_FONT = "Consolas"

_INLINE = re.compile(r"(\*\*.+?\*\*|\*[^*\n]+?\*|`[^`\n]+?`)")


def add_runs(par, text):
    for part in _INLINE.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            r = par.add_run(part[2:-2])
            r.bold = True
        elif part.startswith("`") and part.endswith("`"):
            r = par.add_run(part[1:-1])
            r.font.name = CODE_BG_FONT
            r.font.size = Pt(9.5)
        elif part.startswith("*") and part.endswith("*"):
            r = par.add_run(part[1:-1])
            r.italic = True
        else:
            par.add_run(part)


def main(src, dst):
    lines = open(src, encoding="utf-8").read().splitlines()
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)

    i = 0
    while i < len(lines):
        line = lines[i]

        if line.startswith("```"):                     # code block
            i += 1
            block = []
            while i < len(lines) and not lines[i].startswith("```"):
                block.append(lines[i])
                i += 1
            i += 1
            p = doc.add_paragraph()
            r = p.add_run("\n".join(block))
            r.font.name = CODE_BG_FONT
            r.font.size = Pt(8.5)
            continue

        if line.startswith("|") and i + 1 < len(lines) and \
                re.match(r"^\|[\s\-|:]+\|?$", lines[i + 1]):  # table
            rows = []
            while i < len(lines) and lines[i].startswith("|"):
                if not re.match(r"^\|[\s\-|:]+\|?$", lines[i]):
                    rows.append([c.strip() for c in lines[i].strip("|").split("|")])
                i += 1
            ncols = max(len(r_) for r_ in rows)
            t = doc.add_table(rows=len(rows), cols=ncols)
            t.style = "Light Grid Accent 1"
            t.alignment = WD_TABLE_ALIGNMENT.CENTER
            for ri, row in enumerate(rows):
                for ci in range(ncols):
                    cell = t.cell(ri, ci)
                    cell.paragraphs[0].text = ""
                    add_runs(cell.paragraphs[0], row[ci] if ci < len(row) else "")
                    for par in cell.paragraphs:
                        for run in par.runs:
                            run.font.size = Pt(9)
                            if ri == 0:
                                run.bold = True
            doc.add_paragraph()
            continue

        m = re.match(r"^(#{1,4})\s+(.*)", line)
        if m:                                          # heading
            level = len(m.group(1))
            h = doc.add_heading("", level=min(level, 3))
            add_runs(h, m.group(2))
            for run in h.runs:
                run.font.color.rgb = ACCENT
            i += 1
            continue

        if re.match(r"^\s*([-*])\s+", line):           # bullet list
            while i < len(lines) and re.match(r"^\s*([-*])\s+", lines[i]):
                item = re.sub(r"^\s*[-*]\s+", "", lines[i])
                # join wrapped continuation lines
                while i + 1 < len(lines) and lines[i + 1].startswith("  ") and \
                        not re.match(r"^\s*[-*]\s+", lines[i + 1]):
                    i += 1
                    item += " " + lines[i].strip()
                p = doc.add_paragraph(style="List Bullet")
                add_runs(p, item)
                i += 1
            continue

        if re.match(r"^\s*\d+\.\s+", line):            # numbered list
            while i < len(lines) and re.match(r"^\s*\d+\.\s+", lines[i]):
                item = re.sub(r"^\s*\d+\.\s+", "", lines[i])
                while i + 1 < len(lines) and lines[i + 1].startswith("   ") and \
                        not re.match(r"^\s*\d+\.\s+", lines[i + 1]):
                    i += 1
                    item += " " + lines[i].strip()
                p = doc.add_paragraph(style="List Number")
                add_runs(p, item)
                i += 1
            continue

        if re.match(r"^-{3,}$", line):                 # hr
            i += 1
            continue

        if not line.strip():
            i += 1
            continue

        # paragraph: join soft-wrapped lines until blank/structural
        para = [line]
        while i + 1 < len(lines) and lines[i + 1].strip() and \
                not re.match(r"^(#{1,4}\s|\||```|-{3,}$|\s*[-*]\s|\s*\d+\.\s)", lines[i + 1]):
            i += 1
            para.append(lines[i].strip())
        p = doc.add_paragraph()
        add_runs(p, " ".join(para))
        if para[0].startswith("**Technical report"):
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        i += 1

    doc.save(dst)
    print(f"wrote {dst}")


if __name__ == "__main__":
    main(sys.argv[1], sys.argv[2])
